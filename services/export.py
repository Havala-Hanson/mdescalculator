# services/export.py

import io
import math
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ------------------------------------------------------------
# Kraft (2019) effect-size thresholds for education research
# ------------------------------------------------------------
def classify_effect_size_kraft(d):
    if d < 0.05:
        return "small"
    elif d < 0.20:
        return "medium"
    else:
        return "large"


# ------------------------------------------------------------
# Percentile-shift interpretation for continuous outcomes
# ------------------------------------------------------------
def percentile_shift(d):
    """
    Convert an effect size into a percentile shift using the normal CDF.
    Returns (new_percentile, shift).
    """
    phi = 0.5 * (1 + math.erf(d / math.sqrt(2)))
    shift = (phi - 0.5) * 100
    return phi * 100, shift


# ------------------------------------------------------------
# Styling constants
# ------------------------------------------------------------
_NAVY  = RGBColor(0x1F, 0x38, 0x64)   # #1F3864 – document title
_BLUE  = RGBColor(0x2E, 0x75, 0xB6)   # #2E75B6 – section headers / table header fill
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)   # #FFFFFF – table header text
_IN_BG = RGBColor(0xD9, 0xE1, 0xF2)   # #D9E1F2 – inputs table header fill
_GRAY  = RGBColor(0x59, 0x59, 0x59)   # #595959 – footer / captions


# ------------------------------------------------------------
# Styling helpers
# ------------------------------------------------------------
def _set_cell_bg(cell, rgb: RGBColor):
    """Fill a table cell with a solid background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def _style_run(run, size_pt, bold=False, italic=False,
               color: Optional[RGBColor] = None, font_name='Calibri'):
    run.font.name   = font_name
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_section_heading(doc, text):
    """Bold 14pt Calibri in #2E75B6, with spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    _style_run(run, 14, bold=True, color=_BLUE)
    return p


def _add_body_para(doc, text):
    """11pt Calibri, justified, 1.15 line spacing, 6pt after."""
    p   = doc.add_paragraph(text)
    fmt = p.paragraph_format
    fmt.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.line_spacing = 1.15
    fmt.space_after  = Pt(6)
    for run in p.runs:
        _style_run(run, 11)
    return p


def _add_horizontal_rule(doc):
    """Thin bottom-border paragraph used as a visual divider."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _remove_outer_borders(table):
    """Remove outer four borders of a table, keep inner gridlines."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBdr = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBdr.append(el)
    for side in ('insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:color'), 'auto')
        tblBdr.append(el)
    tblPr.append(tblBdr)


def _make_two_col_table(doc, rows_data,
                        header_labels=('Parameter', 'Selected Value'),
                        header_bg=None):
    """
    Build a styled 2-column table with a coloured header row.

    Parameters
    ----------
    rows_data     : list of (label, value) tuples
    header_labels : column header strings
    header_bg     : RGBColor for header fill; defaults to _BLUE
    """
    bg    = header_bg if header_bg is not None else _BLUE
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for cell, label in zip(hdr, header_labels):
        cell.text = ''
        run = cell.paragraphs[0].add_run(label)
        _style_run(run, 10, bold=True, color=_WHITE)
        _set_cell_bg(cell, bg)

    for label, value in rows_data:
        row   = table.add_row().cells
        run_l = row[0].paragraphs[0].add_run(
            str(label) if label is not None else 'Not provided.'
        )
        run_v = row[1].paragraphs[0].add_run(
            str(value) if value is not None else 'Not provided.'
        )
        _style_run(run_l, 10)
        _style_run(run_v, 10)

    _remove_outer_borders(table)
    return table


def _add_page_number(paragraph):
    """Append a right-aligned 'Page X of Y' field to a paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _fld(text):
        run = paragraph.add_run()
        _style_run(run, 8, italic=True, color=_GRAY)
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = text
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)

    run_pre = paragraph.add_run('Page ')
    _style_run(run_pre, 8, italic=True, color=_GRAY)
    _fld('PAGE')
    run_of = paragraph.add_run(' of ')
    _style_run(run_of, 8, italic=True, color=_GRAY)
    _fld('NUMPAGES')


def _add_footer_top_border(footer):
    """Draw a thin top border on the first footer paragraph."""
    if not footer.paragraphs:
        return
    para = footer.paragraphs[0]
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top  = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '4')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), '595959')
    pBdr.append(top)
    pPr.append(pBdr)


# ------------------------------------------------------------
# Main export function – returns io.BytesIO for Streamlit
# ------------------------------------------------------------

def generate_docx(title: str, inputs: dict, results, narrative: str,
                  calc_narrative: str, metadata: Optional[dict] = None,
                  citations: Optional[list] = None) -> io.BytesIO:
    """
    Build and return a fully-styled MDES .docx report as an in-memory buffer.

    Parameters
    ----------
    title          : descriptive report title (e.g. design name)
    inputs         : ordered dict of user-selected calculation inputs
    results        : result object with attributes mdes, se, df, design_effect,
                     effective_n, total_n; and optionally mdes_pct_points
    narrative      : pre-built interpretive narrative string
    calc_narrative : pre-built methodology / calculation narrative string
    metadata       : reserved for future use
    citations      : optional list of APA citation strings for the footer
    """
    doc = Document()

    # ── Page setup ───────────────────────────────────────────────────────────
    section               = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # ── Title block ──────────────────────────────────────────────────────────
    title_para           = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run            = title_para.add_run(title)
    _style_run(title_run, 20, bold=True, color=_NAVY)

    subtitle_para           = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y %H:%M')}"
    )
    _style_run(sub_run, 11, color=_GRAY)

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # ── Calculation Inputs ───────────────────────────────────────────────────
    _add_section_heading(doc, 'Calculation Inputs')
    _make_two_col_table(
        doc,
        list(inputs.items()),
        header_labels=('Parameter', 'Selected Value'),
        header_bg=_IN_BG,
    )
    doc.add_paragraph()

    # ── Results ──────────────────────────────────────────────────────────────
    _add_section_heading(doc, 'Results')

    outcome_type = inputs.get('outcome_type', 'continuous')
    if outcome_type == 'binary':
        mdes_label = 'MDES (percentage points)'
    else:
        mdes_label = 'MDES (standardized)'

    result_rows = [
        (mdes_label,              results.mdes),
        ('Standard error',        results.se),
        ('Degrees of freedom',    results.df),
        ('Two-tailed or one-tailed', inputs.get('two_tailed', True) and 'Two-tailed' or 'One-tailed'),
        ('Design effect',         results.design_effect),
        ('Effective sample size', results.effective_n),
        ('Total sample size',     results.total_n),
    ]
    if getattr(results, 'mdes_pct_points', None) is not None and outcome_type != 'binary':
        result_rows.insert(1, ('MDES (percentage points)', results.mdes_pct_points))
    if getattr(results, 'mdes_standardized', None) is not None and outcome_type != 'binary':
        result_rows.insert(1, ('MDES (standardized)', results.mdes_standardized))

    res_table = doc.add_table(rows=1, cols=2)
    res_table.style = 'Table Grid'
    res_hdr = res_table.rows[0].cells
    for cell, label in zip(res_hdr, ('Metric', 'Value')):
        cell.text = ''
        run = cell.paragraphs[0].add_run(label)
        _style_run(run, 10, bold=True, color=_WHITE)
        _set_cell_bg(cell, _BLUE)

    for i, (metric, value) in enumerate(result_rows):
        row     = res_table.add_row().cells
        is_last = (i == len(result_rows) - 1)
        val_str = str(value) if value is not None else 'Not provided.'
        run_m   = row[0].paragraphs[0].add_run(str(metric))
        run_v   = row[1].paragraphs[0].add_run(val_str)
        _style_run(run_m, 10, bold=is_last)
        _style_run(run_v, 10, bold=is_last)
        row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    _remove_outer_borders(res_table)
    doc.add_paragraph()

    # ── Interpretive Narrative ───────────────────────────────────────────────
    _add_section_heading(doc, 'Interpretation')
    _add_body_para(doc, narrative if narrative else 'Not provided.')

    # ── Methodology & Calculations ───────────────────────────────────────────
    _add_section_heading(doc, 'Methodology & Calculations')
    _add_body_para(doc, calc_narrative if calc_narrative else 'Not provided.')

    # ── Footer (all pages) ───────────────────────────────────────────────────
    default_citations = [
        (
            'Dong, N., & Maynard, R. (2013). PowerUp!: A tool for calculating minimum '
            'detectable effect sizes and minimum required sample sizes for experimental '
            'and quasi-experimental design studies. Journal of Research on Educational '
            'Effectiveness, 6(1), 24\u201367. '
            'https://doi.org/10.1080/19345747.2012.673143'
        ),
        (
            'Kraft, M. A. (2020). Interpreting effect sizes of education interventions. '
            'Educational Researcher, 49(4), 241\u2013253. '
            'https://doi.org/10.3102/0013189X20912798'
        ),
    ]
    cite_list = citations if citations else default_citations

    footer = section.footer
    footer.is_linked_to_previous = False
    _add_footer_top_border(footer)

    # Left-aligned APA citations
    cite_para           = footer.paragraphs[0]
    cite_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for i, cite in enumerate(cite_list):
        run = cite_para.add_run(('' if i == 0 else '\n') + cite)
        _style_run(run, 8, italic=True, color=_GRAY)

    # Centred branding line
    branding_para           = footer.add_paragraph()
    branding_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    branding_run            = branding_para.add_run('Created with Precision MDES  v1.0.0')
    _style_run(branding_run, 8, italic=True, color=_GRAY)

    # Right-aligned page number
    page_para = footer.add_paragraph()
    _add_page_number(page_para)

    # ── Return buffer ────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer